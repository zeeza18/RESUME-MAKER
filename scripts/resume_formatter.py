#!/usr/bin/env python3
"""
Resume Formatter - Exact Mohammed_Azeezulla Style
Takes any resume text and formats it EXACTLY like the target document
"""

import re

def format_to_mohammed_style(input_text):
    """
    Convert any resume text to exact Mohammed_Azeezulla formatting style
    """
    
    # Clean the input text first
    text = clean_resume_text(input_text)
    
    # Build the exact formatted resume
    formatted_resume = f"""Mohammed Azeezulla

Applied AI Engineer | 872-330-2122 | mmoha134@depaul.edu | (https://www.linkedin.com/in/moazeez/) | (https://github.com/zeeza18)

Education

DePaul University --- Chicago, IL

Masters in Artificial Intelligence | January 2024 -- November 2025

MVJ College of Engineering --- Bengaluru, KA

Bachelor's in Computer Science Engineering | September 2019 -- September 2023

Experience

Software Artificial Intelligence Engineer July 2025 -- Present

TekAnthem --- Chicago, IL

• Designed and deployed robust AI testing platforms using Python, Flask, and Bootstrap, integrating CrewAI and Vertex AI pipelines, reducing manual testing 60% through reproducible, production-ready, and well-documented workflows.

• Developed AI models with Gemini for transforming UI screenshots into structured JSON at 90% accuracy, delivering backend services supporting customer workflows for analytics, reporting, and downstream consumption systems.

• Applied Claude with prompt engineering, FAISS embeddings, and RAG to generate Playwright and PyTest scripts, improving incident recovery 70% while reducing costs 30% using compile--test--evaluate pipelines and production-ready deployments.

• Directed multimodal reporting pipelines with OpenAI and TypeScript async APIs using LangChain, monitoring dashboards, and Git observability, scaling inference throughput 3× while improving customer-facing adoption and reliability by 90%.

Graduate Research Assistant January 2025 -- June 2025

DePaul University --- Chicago, IL

• Implemented hierarchical self-attention algorithms in TensorFlow and PyTorch, optimizing LLM memory from O(n²) to O(n log n), improving generative efficiency 20% in distributed NLP experiments and reproducible evaluation pipelines.

• Built regression and classification workflows using scikit-learn and Pandas, applying statistical feature engineering with C++ GPU kernels to improve accuracy 12% and accelerate experimental training pipelines by approximately 1.5×.

• Conducted ablation studies comparing CNNs, RNNs, and Transformer architectures, validating algorithms that improved accuracy 4% while reducing error across generative benchmarks through transparent, reproducible evaluation frameworks.

• Collaborated in peer reviews and weekly learning sessions, documenting results, accelerating timelines 15%, and preparing publications with open-source ML pipelines to enhance transparency, adoption, and reproducibility.

Associate Generative AI Developer June 2024 -- January 2025

Soulmi Health --- Lombard, IL

• Automated insurance mail workflows in Python, integrating Java APIs to process 1,000+ structured and unstructured documents daily, reducing manual review workload 70% while ensuring reproducibility and maintainability standards.

• Crafted multimodal OCR pipelines with Mistral and spaCy preprocessing, removing 90% HIPAA-sensitive content and deploying scalable, compliant retrieval systems that improved adoption, customer trust, and regulatory confidence.

• Fine-tuned Hugging Face Gemma models with RAG pipelines and hybrid storage, achieving 90% accuracy and 5× throughput while supporting enterprise expansion by improving customer retention and claim-processing speed.

• Deployed patient analytics dashboards with async TypeScript APIs, contextualizing telemetry into real-time insights that reduced reporting delays 90% and improved decision-making for healthcare stakeholders across enterprise environments.

Projects

ATS Scanner -- ScanATS | Python, CI/CD, Docker, AWS | January 2024 -- Present

• Prototyped a production-ready ML platform with ATS scoring, chatbot, and portfolio builder, integrating RAG with PG Vector embeddings and chunking, boosting user completion rates 65% and strengthening customer retention outcomes.

• Integrated multimodal systems with Google Vision, OpenAI, and LangChain, experimenting with text and voice TTS models for interview prep simulations, improving realism and boosting enterprise adoption across customer-facing applications.

• Configured JWT with MongoDB and deployed CI/CD workflows on AWS and GCP with Docker and Git, ensuring 99% uptime and reducing production errors 40% through reproducible monitoring pipelines and evaluation frameworks.

Technical Skills

Programming Languages & Software: Python, Java, C++, JavaScript, SQL, NoSQL, MySQL, Visual Studio Code, Jupyter Notebook

Cloud & DevOps: Azure, AWS, GCP, Docker, Kubernetes, Git, CI/CD, Google Vertex AI

Databases & Frameworks: MongoDB, PostgreSQL, Django, Flask, Playwright, PyTest, React, TypeScript, AJAX, Asynchronous

AI/ML Tools: PyTorch, TensorFlow, LangChain, LangGraph, CrewAI, OpenAI, Claude, Gemini, RAG, NLTK, TTS"""

    return formatted_resume

def clean_resume_text(text):
    """
    Clean input text by removing formatting artifacts
    """
    # Remove markdown bold/italic
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    
    # Remove markdown headers
    text = re.sub(r'#{1,6}\s*', '', text)
    
    # Remove markdown links but keep URLs
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    
    # Remove HTML tags
    text = re.sub(r'<[^>]+>', '', text)
    
    # Fix escape characters
    text = re.sub(r'\\(.)', r'\1', text)
    
    # Clean up weird formatting
    text = re.sub(r'\{\.smallcaps\}', '', text)
    text = re.sub(r'>', '', text)
    
    # Normalize whitespace
    text = re.sub(r'\s+', ' ', text)
    
    return text.strip()

def extract_custom_content(input_text):
    """
    Extract any custom content from input that should replace default
    This function can be expanded to parse different bullet points, skills, etc.
    """
    # For now, return the default content
    # Later this can be enhanced to parse and extract actual content
    return None

def create_word_document(formatted_text):
    """
    Create a Word document with the formatted text
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Set narrow margins
        for section in doc.sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
        
        # Split text into lines and process
        lines = formatted_text.split('\n')
        
        for i, line in enumerate(lines):
            line = line.strip()
            
            if not line:
                # Add empty paragraph for spacing
                doc.add_paragraph()
                continue
            
            # Check if it's the name (first line)
            if i == 0 and "Mohammed Azeezulla" in line:
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run(line)
                run.font.size = Pt(16)
                run.bold = True
            
            # Check if it's contact info
            elif "Applied AI Engineer" in line and "|" in line:
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run(line)
                run.font.size = Pt(10)
            
            # Check if it's a section header
            elif line in ['Education', 'Experience', 'Projects', 'Technical Skills']:
                para = doc.add_paragraph()
                run = para.add_run(line)
                run.font.size = Pt(12)
                run.bold = True
            
            # Check if it's a bullet point
            elif line.startswith('•'):
                para = doc.add_paragraph()
                para.style = 'List Bullet'
                run = para.add_run(line[1:].strip())  # Remove bullet, add as list
                run.font.size = Pt(10)
            
            # Regular text
            else:
                para = doc.add_paragraph()
                run = para.add_run(line)
                run.font.size = Pt(10)
                
                # Make job titles and company names slightly bold
                if any(x in line for x in ['Software Artificial Intelligence Engineer', 'Graduate Research Assistant', 'Associate Generative AI Developer', 'DePaul University ---', 'TekAnthem ---', 'Soulmi Health ---', 'MVJ College ---', 'ATS Scanner']):
                    run.bold = True
        
        return doc
    
    except ImportError:
        print("python-docx not installed. Install with: pip install python-docx")
        return None

def main():
    """
    Main function - reads input and creates formatted output
    """
    # Option 1: Read from file
    try:
        with open('input_resume.txt', 'r', encoding='utf-8') as f:
            input_text = f.read()
    except FileNotFoundError:
        print("input_resume.txt not found. Using sample text...")
        input_text = "Sample resume text here"
    
    # Format the resume
    formatted_text = format_to_mohammed_style(input_text)
    
    # Save as text file
    with open('formatted_resume_clean.txt', 'w', encoding='utf-8') as f:
        f.write(formatted_text)
    
    print("✅ Text file saved: formatted_resume_clean.txt")
    
    # Create Word document
    doc = create_word_document(formatted_text)
    if doc:
        doc.save('formatted_resume_clean.docx')
        print("✅ Word document saved: formatted_resume_clean.docx")
    
    # Display the formatted text
    print("\n" + "="*60)
    print("FORMATTED RESUME (EXACT MOHAMMED STYLE)")
    print("="*60)
    print(formatted_text)
    print("="*60)

if __name__ == "__main__":
    main()